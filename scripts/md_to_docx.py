from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

root = Path(r"C:\Users\sweta\OneDrive\Documents\capstone2draft2")
md_path = root / "Multimodal AI Research Assistant using Retrieval-Augmented Generation (RAG) - Final Report.md"
docx_path = root / "Multimodal AI Research Assistant using Retrieval-Augmented Generation (RAG) - Final Report.docx"

lines = md_path.read_text(encoding="utf-8").splitlines()
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)

styles = doc.styles
for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "List Bullet", "List Number"]:
    style = styles[style_name]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
if "Normal" in styles:
    styles["Normal"].font.size = Pt(12)

for raw_line in lines:
    line = raw_line.rstrip()
    if not line:
        doc.add_paragraph("")
        continue
    if line.startswith("# "):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line[2:])
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(20)
        continue
    if line.startswith("## "):
        h = doc.add_paragraph(style="Heading 1")
        r = h.add_run(line[3:])
        r.bold = True
        r.font.name = "Times New Roman"
        continue
    if line.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(line[2:])
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        continue
    if line[:2].isdigit() and line[2:4] == ". ":
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(line[4:])
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        continue
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(line)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

doc.save(docx_path)
print(docx_path)